#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def set_metadata(document: Document) -> None:
    props = document.core_properties
    props.title = "ASMCM Deliverables Template"
    props.subject = "Template for paper, presentation, and single-blind peer review"
    props.author = "OpenAI"
    props.language = "en-US"


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Aptos"

    document.styles["Title"].font.size = Pt(22)
    document.styles["Heading 1"].font.size = Pt(16)
    document.styles["Heading 2"].font.size = Pt(13)
    document.styles["Heading 3"].font.size = Pt(11)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.text = "ASMCM Deliverables Template | "
        add_page_number(footer)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Experiment / Task Group"
    hdr[1].text = "Model / Setup"
    hdr[2].text = "Observation / Status"
    rows = [
        ("Label tasks", "[Model name]", "[Accuracy, notes, or pending]"),
        ("Sequence tasks", "[Model name]", "[Exact match, pitch accuracy, or pending]"),
        ("Explanatory mode", "[Model name]", "[Qualitative comments or pending]"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ASMCM Deliverables Template")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Prepared from the deliverables list dated March 18, 2026\n"
        "Project working focus: A Symbolic Music Benchmark for Music Theory and Computational Musicology"
    )

    document.add_paragraph("")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Student(s): ").bold = True
    meta.add_run("[Your name] / [Teammate name]\n")
    meta.add_run("Course / Instructor: ").bold = True
    meta.add_run("[Fill in]\n")
    meta.add_run("Submission date: ").bold = True
    meta.add_run("[Fill in]")

    document.add_paragraph("")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "This template already includes a suggested framing, sample text, and section prompts. "
        "Replace bracketed placeholders with your final content."
    )


def add_paper_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("1. Written Paper (Main Submission)", level=1)

    document.add_heading("Working Title", level=2)
    document.add_paragraph(
        "A Symbolic Music Benchmark for Music Theory and Computational Musicology"
    )

    document.add_heading("Abstract", level=2)
    document.add_paragraph(
        "This paper investigates whether language-model-based systems can solve "
        "symbolic tasks grounded in music-theory and computational musicology. "
        "We frame the project as a multi-task benchmark that combines label classification "
        "tasks such as interval, chord, harmonic-function, and voice-leading analysis with "
        "sequence transformation tasks such as transposition, melodic inversion, retrograde, "
        "and rhythmic scaling. The benchmark is designed for automatic formal evaluation while "
        "also supporting an optional explanation layer for qualitative analytical discussion. "
        "This draft should later summarize the final system, dataset construction, evaluation "
        "pipeline, and key empirical findings in no more than about 200 words."
    )

    document.add_heading("Introduction", level=2)
    document.add_paragraph(
        "Suggested opening angle: music-theory tasks are often discussed qualitatively, "
        "but benchmark design requires operational definitions, stable symbolic representations, "
        "and measurable outputs. This project turns those analytical concepts into reproducible tasks."
    )
    document.add_heading("Suggested research questions", level=3)
    add_numbered(
        document,
        [
            "To what extent can a language model recognize, transform, and explain symbolic musical relationships?",
            "How well do standard benchmark metrics capture musically meaningful correctness?",
            "What is gained by separating strict benchmark outputs from optional analytical explanations?",
        ],
    )

    document.add_heading("Related Work", level=2)
    add_bullets(
        document,
        [
            "Music representation learning and symbolic music modeling",
            "Music-theory-oriented benchmarks or rule-based music-analysis systems",
            "LLM evaluation for structured reasoning tasks",
            "Computational musicology and symbolic analytical frameworks",
        ],
    )
    document.add_paragraph(
        "Writing prompt: explain how your benchmark differs from prior symbolic music tasks "
        "by combining formal automatic evaluation with a musicology-facing explanation layer."
    )

    document.add_heading("Methodology", level=2)
    document.add_paragraph(
        "Recommended framing for this section: the system is built as a symbolic benchmark "
        "pipeline with rule-based case generation, tokenizer views, model-ready case packaging, "
        "canonical prediction fields, and unified evaluation."
    )
    add_bullets(
        document,
        [
            "Task families: label classification and note-sequence transformation",
            "Current V1 scope: major-mode-only harmonic function, two-voice voice-leading, reduced harmonic vocabulary",
            "Standard mode: prediction_label / prediction_notes with strict automatic evaluation",
            "Explanatory mode: optional prediction_explanation for paper-oriented qualitative analysis",
            "Evaluation outputs: overall, by-task, by-tokenizer, and annotated review files",
        ],
    )

    document.add_heading("Results", level=2)
    document.add_paragraph(
        "Use this section to report benchmark runs, early observations, or even clearly marked expected outcomes if final experiments are not yet complete."
    )
    add_table(document)

    document.add_heading("Discussion / Conclusion", level=2)
    add_bullets(
        document,
        [
            "Interpret what the benchmark reveals about symbolic musical reasoning",
            "Explain the meaning of the V1 scope and planned extensions",
            "Discuss the tension between structural correctness and analytical explanation quality",
            "State limitations and future work clearly",
        ],
    )

    document.add_heading("References", level=2)
    document.add_paragraph(
        "[Insert properly formatted references here. Use a consistent citation style.]"
    )


def add_presentation_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("2. Paper Presentation", level=1)
    document.add_paragraph(
        "Suggested presentation goal: explain the problem, the benchmark design, the current V1 scope, and the value of the project as both a software artifact and a paper proposal."
    )
    document.add_heading("Suggested slide flow", level=2)
    add_numbered(
        document,
        [
            "Title and project framing",
            "Problem statement and motivation",
            "Benchmark task families and outputs",
            "Current V1 scope",
            "Planned extensions and not yet final theoretical coverage",
            "Architecture and workflow",
            "Evaluation strategy",
            "Musicology layer: standard vs explanatory mode",
            "Current status, limitations, and next steps",
        ],
    )
    document.add_heading("Presentation checklist", level=2)
    add_bullets(
        document,
        [
            "State the problem clearly in the first 1-2 slides",
            "Show why the benchmark is musically meaningful, not only technically structured",
            "Use at least one architecture or workflow diagram",
            "Keep the explanation of V1 scope honest and explicit",
            "End with contributions, limitations, and next steps",
        ],
    )
    document.add_paragraph(
        "Practical note: the current project deck can be aligned with this section, but this template lets you rewrite the oral narrative before final slides are frozen."
    )


def add_review_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("3. Single-Blind Peer Review", level=1)
    document.add_paragraph(
        "Use this as a review-report scaffold for another paper. Keep the tone professional, specific, and constructive."
    )

    review_blocks = [
        ("Paper summary", "[Briefly summarize the paper's main goal, method, and contribution.]"),
        ("Strengths", "[List the paper's strongest ideas, methods, or presentation choices.]"),
        ("Weaknesses", "[List the most important limitations or unclear aspects.]"),
        ("Comments on title and abstract", "[Assess relevance, clarity, and alignment with the actual content.]"),
        ("Comments on related work", "[Assess coverage, positioning, and adequacy of citations.]"),
        ("Comments on methodology and evaluation", "[Assess soundness, reproducibility, and whether the evaluation supports the claims.]"),
        ("Novelty and contribution", "[Assess what is genuinely new or useful in the paper.]"),
        ("Comments on presentation", "[Assess writing clarity, structure, figures, and readability.]"),
        ("Overall recommendation", "[Accept / weak accept / borderline / weak reject / reject, with justification.]"),
    ]
    for title, prompt in review_blocks:
        document.add_heading(title, level=2)
        document.add_paragraph(prompt)


def add_minimum_requirements(document: Document) -> None:
    document.add_page_break()
    document.add_heading("4. Minimal Requirements Checklist", level=1)
    add_bullets(
        document,
        [
            "A real draft of the written paper, not only notes",
            "A clear research direction and section structure",
            "Initial preparation for the presentation phase",
            "Initial preparation for the peer-review phase",
        ],
    )
    document.add_paragraph(
        "Suggested immediate milestone for this project: finish a coherent draft with title, abstract, introduction, methodology, and a provisional results/discussion structure, even if some final experiments are still pending."
    )


def build_document(output_path: Path) -> None:
    document = Document()
    set_margins(document)
    set_metadata(document)
    style_document(document)
    add_footer(document)

    add_title_page(document)
    add_paper_section(document)
    add_presentation_section(document)
    add_review_section(document)
    add_minimum_requirements(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    out = SUPPORTING_DOCS_DIR / "docx" / "ASMCM-deliverables-template.docx"
    build_document(out)
    print(f"wrote: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
SUPPORTING_DOCS_DIR = Path(__file__).resolve().parents[1]
