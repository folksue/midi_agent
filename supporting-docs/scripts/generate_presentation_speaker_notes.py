#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SUPPORTING_DOCS_DIR = Path(__file__).resolve().parents[1]
MD_OUT = SUPPORTING_DOCS_DIR / "notes" / "ASMCM-presentation-speaker-notes.md"
DOCX_OUT = SUPPORTING_DOCS_DIR / "docx" / "ASMCM-presentation-speaker-notes.docx"


SECTIONS = [
    (
        "Quick Assessment",
        [
            "The presentation makes sense overall and already communicates a coherent project direction.",
            "Its strongest points are the benchmark framing, the clear separation between standard and explanatory modes, and the explicit discussion of V1 scope.",
            "Its weakest point is alignment with the presentation deliverable: it currently emphasizes repository structure more than problem statement, expected outcomes, and final takeaways.",
        ],
    ),
    (
        "Recommended Changes Before Presenting",
        [
            "Make the problem statement more explicit on Slide 3. Right now the motivation is good, but the central research problem should be stated in one short sentence.",
            "Add or strengthen an expected outcomes slide. The deliverables explicitly ask for key results or expected outcomes, and this is still underdeveloped in the current deck.",
            "Turn Slide 15 into a stronger conclusion. It should summarize the contribution, expected findings, and next steps in a clearer closing message.",
            "If presentation time is short, move Slides 12, 13, and 14 to backup or appendix. They are useful, but they are more technical than necessary for a first proposal talk.",
            "Fix visible wording issues: 'Taks' on the agenda slide, 'Non tech user' on the use-case slide, 'can now be as both' on the last slide, and 'Enhacne' in the next-steps box.",
            "Consider merging Slides 7 and 8 if you need a tighter story. Architecture and workflow are both useful, but together they may be more detail than required for the deliverable.",
        ],
    ),
    (
        "Deliverables Alignment",
        [
            "Problem statement and motivation: mostly covered by Slide 3, but the problem statement should be phrased more directly.",
            "Overview of methodology: covered across Slides 4, 7, and 8.",
            "Key results or expected outcomes: partially implied, but not yet clearly presented. This is the main gap relative to the deliverable.",
            "Conclusions: partially present in Slide 15, but still too diffuse. The ending should be sharper.",
        ],
    ),
    (
        "Suggested Core Deck For A Short Talk",
        [
            "Slide 1: Title and framing",
            "Slide 3: Motivation and problem statement",
            "Slide 4: Benchmark tasks and outputs",
            "Slide 5: Current V1 scope",
            "Slide 7 or 8: Method overview",
            "Slide 10: Evaluation strategy",
            "New or revised slide: Expected outcomes / current status",
            "Slide 11: Why the musicology layer matters",
            "Slide 15: Conclusion and next steps",
        ],
    ),
]


SLIDE_NOTES = [
    {
        "slide": "Slide 1 - Title",
        "purpose": "Open the talk and frame the project in one sentence.",
        "notes": (
            "Start with the simplest possible explanation: this project builds a symbolic music benchmark "
            "to test whether AI models can solve music-theory tasks in a formally correct and analytically meaningful way. "
            "Emphasize that the project is positioned between music theory, computational musicology, and model evaluation."
        ),
        "transition": "Then move into why a benchmark like this is needed."
    },
    {
        "slide": "Slide 2 - Agenda",
        "purpose": "Give the audience a roadmap.",
        "notes": (
            "Keep this short. Say that you will first explain the motivation, then the task design, then the current scope, "
            "then the technical workflow, and finally why this matters as a paper contribution."
        ),
        "transition": "Now that the structure is clear, explain the motivation."
    },
    {
        "slide": "Slide 3 - Motivation",
        "purpose": "State the problem and justify the project.",
        "notes": (
            "This is the most important early slide. Explain that symbolic musical reasoning is a good evaluation target because "
            "it is structured, rule-based, and musically meaningful. Add an explicit problem sentence such as: "
            "'The problem is that we still lack a clean benchmark for evaluating whether a model understands symbolic music relationships rather than only generating plausible tokens.'"
        ),
        "transition": "After motivating the benchmark, describe what the benchmark actually contains."
    },
    {
        "slide": "Slide 4 - Tasks",
        "purpose": "Explain the task families and outputs.",
        "notes": (
            "Present the two big categories clearly: classification tasks and transformation tasks. "
            "Tell the audience that this design allows you to evaluate both recognition and manipulation of symbolic musical structures. "
            "Mention that the outputs are intentionally simple so they can be scored automatically."
        ),
        "transition": "Once the tasks are clear, explain what the current benchmark does and does not try to cover."
    },
    {
        "slide": "Slide 5 - Current V1 Scope",
        "purpose": "Explain the current boundary of the benchmark.",
        "notes": (
            "Make it clear that these are not arbitrary shortcuts but deliberate V1 constraints. "
            "Say that the current version focuses on a narrow but defensible slice of symbolic tonal reasoning so that the benchmark stays testable and reproducible."
        ),
        "transition": "After defining the current scope, clarify how this could expand later."
    },
    {
        "slide": "Slide 6 - Planned Extensions",
        "purpose": "Show that the project has a growth path.",
        "notes": (
            "Use this slide to show that the benchmark is not claiming complete theoretical coverage. "
            "Say that the current version is scientifically usable now, but later versions can broaden harmonic coverage, voice-leading complexity, and the explanation layer."
        ),
        "transition": "Now that the conceptual scope is clear, show how the system works."
    },
    {
        "slide": "Slide 7 - Architecture",
        "purpose": "Explain the benchmark pipeline at a high level.",
        "notes": (
            "Do not over-explain file names unless the audience is highly technical. "
            "The key message is that the benchmark has a rule-based generation stage, a case-building stage, a model-inference stage, and an evaluation stage."
        ),
        "transition": "After the architecture, move to the user-facing workflow."
    },
    {
        "slide": "Slide 8 - Workflow",
        "purpose": "Show the practical run sequence.",
        "notes": (
            "This is the operational version of the previous slide. Explain that a user can generate cases, export views, build model-ready cases, run a model, and automatically evaluate the outputs. "
            "Use this slide to emphasize usability for non-expert users."
        ),
        "transition": "Next, show how predictions are represented."
    },
    {
        "slide": "Slide 9 - Prediction Schema",
        "purpose": "Explain how answers are encoded.",
        "notes": (
            "Keep this brief unless your audience is technical. The important message is that classification tasks return a label, "
            "sequence tasks return symbolic notes, and everything is kept in a canonical format for automatic evaluation."
        ),
        "transition": "After the output format, explain how you score the system."
    },
    {
        "slide": "Slide 10 - Evaluation",
        "purpose": "Explain how correctness is measured.",
        "notes": (
            "This slide is important because it shows that the benchmark is not vague. "
            "Say that different tasks use different metrics, but all are formally defined and automatically computable. "
            "If you revise this deck, this is also a good place to mention expected outcomes or the types of findings you want to obtain."
        ),
        "transition": "Then explain why the project is not only technical but also musicological."
    },
    {
        "slide": "Slide 11 - Musicology Layer",
        "purpose": "Justify the explanatory dimension.",
        "notes": (
            "Explain that the benchmark keeps a strict standard mode for scoring, but optionally allows explanatory outputs for qualitative analysis. "
            "This is the slide that best connects the benchmark to computational musicology, because it shows how formal correctness and analytical interpretation can be compared."
        ),
        "transition": "From here, the remaining diagram slides can be presented briefly or treated as backup."
    },
    {
        "slide": "Slide 12 - Use-Case View",
        "purpose": "Show who the benchmark is for.",
        "notes": (
            "Use this only if time allows. Explain that the benchmark supports different audiences: a general user who wants to run tests, "
            "a research user who wants paper examples, and a developer who wants model comparisons."
        ),
        "transition": "If you continue, move to the runtime and component diagrams."
    },
    {
        "slide": "Slide 13 - Sequence View",
        "purpose": "Illustrate runtime order.",
        "notes": (
            "Present this only if the audience wants technical detail. The goal is simply to show the sequence from generation to evaluation and review."
        ),
        "transition": "Then move to the component-level summary."
    },
    {
        "slide": "Slide 14 - Component View",
        "purpose": "Summarize system modules.",
        "notes": (
            "Again, this is more of a technical backup slide. Explain that the benchmark is modular: theory logic, representation, execution, and analysis."
        ),
        "transition": "Finally, close with the paper framing and next steps."
    },
    {
        "slide": "Slide 15 - Paper Framing and Next Steps",
        "purpose": "Close the talk with the contribution and next steps.",
        "notes": (
            "Strengthen this ending when you revise the slide. The closing message should be: this is already a usable symbolic benchmark, "
            "it supports a real paper contribution, and the next step is to run experiments and turn the current design into empirical evidence."
        ),
        "transition": "End by restating the main contribution in one sentence."
    },
]


def add_page_number(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph.add_run()._r.append(fld)


def init_doc() -> Document:
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    document.styles["Title"].font.name = "Aptos"
    document.styles["Title"].font.size = Pt(21)
    document.styles["Heading 1"].font.name = "Aptos"
    document.styles["Heading 2"].font.name = "Aptos"

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "ASMCM Presentation Notes | "
    add_page_number(footer)
    return document


def build_markdown() -> str:
    lines: list[str] = []
    lines.append("# ASMCM Presentation Speaker Notes and Review")
    lines.append("")
    lines.append("This document reviews the current presentation against the ASMCM deliverables and provides slide-by-slide speaker notes.")
    lines.append("")

    for heading, items in SECTIONS:
        lines.append(f"## {heading}")
        lines.append("")
        for item in items:
            lines.append(f"- {item}")
        lines.append("")

    lines.append("## Slide-by-Slide Speaker Notes")
    lines.append("")
    for item in SLIDE_NOTES:
        lines.append(f"### {item['slide']}")
        lines.append("")
        lines.append(f"**Purpose:** {item['purpose']}")
        lines.append("")
        lines.append(f"**What to say:** {item['notes']}")
        lines.append("")
        lines.append(f"**Transition:** {item['transition']}")
        lines.append("")
    return "\n".join(lines)


def build_docx() -> Document:
    doc = init_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ASMCM Presentation Speaker Notes and Review")
    run.bold = True
    run.font.size = Pt(21)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Based on the current presentation deck and the ASMCM deliverables list"
    )

    doc.add_paragraph("")

    for heading, items in SECTIONS:
        doc.add_heading(heading, level=1)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Slide-by-Slide Speaker Notes", level=1)
    for item in SLIDE_NOTES:
        doc.add_heading(item["slide"], level=2)
        p1 = doc.add_paragraph()
        p1.add_run("Purpose: ").bold = True
        p1.add_run(item["purpose"])
        p2 = doc.add_paragraph()
        p2.add_run("What to say: ").bold = True
        p2.add_run(item["notes"])
        p3 = doc.add_paragraph()
        p3.add_run("Transition: ").bold = True
        p3.add_run(item["transition"])

    return doc


def main() -> int:
    MD_OUT.parent.mkdir(parents=True, exist_ok=True)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)

    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    doc = build_docx()
    doc.save(DOCX_OUT)

    print(f"wrote: {MD_OUT}")
    print(f"wrote: {DOCX_OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
