#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PUBLICATION_DIR = Path(__file__).resolve().parents[1]
OUT_PATH = PUBLICATION_DIR / "ASMCM-scientific-paper-template.docx"


def set_page_layout(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def set_metadata(document: Document) -> None:
    props = document.core_properties
    props.title = "ASMCM Scientific Paper Template"
    props.subject = "Scientific paper template for the ASMCM project"
    props.author = "OpenAI"
    props.language = "en-US"


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = "Aptos"

    document.styles["Title"].font.size = Pt(22)
    document.styles["Heading 1"].font.size = Pt(16)
    document.styles["Heading 2"].font.size = Pt(13)
    document.styles["Heading 3"].font.size = Pt(11)


def add_page_number(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph.add_run()._r.append(fld)


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.text = "ASMCM Scientific Paper Template | "
        add_page_number(footer)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_caption(document: Document, label: str, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}. {text}")
    run.italic = True
    run.font.size = Pt(10)


def add_experiment_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Task Group"
    hdr[1].text = "Model / System"
    hdr[2].text = "Metric(s)"
    hdr[3].text = "Main Observation"
    rows = [
        ("Label classification", "[Model name]", "[Accuracy / Macro F1]", "[Insert main finding]"),
        ("Sequence transformation", "[Model name]", "[Exact match / Pitch accuracy / Duration accuracy]", "[Insert main finding]"),
        ("Explanatory mode", "[Model name]", "[Qualitative review]", "[Insert analytical comment]"),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def add_repro_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Status / Note"
    rows = [
        ("Repository / code availability", "[Insert link or note]"),
        ("Dataset / benchmark cases", "[Insert location or note]"),
        ("Model configuration", "[Insert model names, prompts, and setup]"),
        ("Evaluation scripts", "[Insert script names or paths]"),
        ("Limitations of current V1 scope", "[Insert concise note]"),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Scientific Paper Template")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Project framing: A Symbolic Music Benchmark for Music Theory and Computational Musicology"
    )

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Authors: ").bold = True
    meta.add_run("[Your name], [Teammate name]\n")
    meta.add_run("Affiliation: ").bold = True
    meta.add_run("[University / Department]\n")
    meta.add_run("Date: ").bold = True
    meta.add_run("[Fill in]\n")
    meta.add_run("Course / Seminar: ").bold = True
    meta.add_run("[Fill in]")

    document.add_paragraph("")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "This is a full article template in English. It includes section structure, "
        "starter prose, writing prompts, and placeholder tables aligned with the current project."
    )


def add_front_matter(document: Document) -> None:
    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "This paper presents a symbolic benchmark for music theory and computational "
        "musicology. The benchmark is organized as a multi-task evaluation suite that combines "
        "label-classification tasks such as interval identification, chord recognition, harmonic "
        "function classification, and voice-leading violation detection with note-sequence "
        "transformation tasks such as transposition, melodic inversion, retrograde, and rhythmic "
        "scaling. The benchmark is designed to support automatic formal evaluation while also "
        "allowing an optional analytical explanation layer for qualitative interpretation. This "
        "template should later be updated with the final system description, experimental setup, "
        "and empirical findings in approximately 150-200 words."
    )

    document.add_heading("Keywords", level=1)
    document.add_paragraph(
        "symbolic music, music theory, computational musicology, benchmark design, language models, evaluation"
    )


def add_introduction(document: Document) -> None:
    document.add_heading("1. Introduction", level=1)
    document.add_paragraph(
        "Music-theory tasks are often discussed in analytical terms, but benchmark construction "
        "requires operational definitions, symbolic representations, and measurable outputs. This "
        "project addresses that gap by defining a benchmark that turns core music-theory concepts "
        "into reproducible symbolic tasks. The central motivation is not only to test whether a "
        "model can produce formally correct answers, but also to study whether those answers can be "
        "interpreted as musically meaningful reasoning."
    )
    document.add_paragraph(
        "Suggested introduction structure: begin with the problem context, explain why symbolic "
        "music is a good evaluation medium, and then position the benchmark as both a technical "
        "artifact and a musicology-facing research contribution."
    )

    document.add_heading("1.1 Research Problem", level=2)
    document.add_paragraph(
        "The research problem can be stated as follows: how can we evaluate symbolic musical "
        "reasoning in a way that is structurally precise, automatically testable, and still "
        "connected to categories used in music theory and computational musicology?"
    )

    document.add_heading("1.2 Research Questions", level=2)
    add_numbered(
        document,
        [
            "To what extent can a model recognize basic symbolic musical relationships such as intervals, chords, and harmonic functions?",
            "To what extent can a model correctly perform symbolic musical transformations such as transposition, inversion, retrograde, and rhythmic scaling?",
            "What is gained by separating strict benchmark outputs from optional analytical explanations?",
            "What do benchmark results reveal about the limits of current models in basic musicological categories?",
        ],
    )

    document.add_heading("1.3 Main Contributions", level=2)
    add_bullets(
        document,
        [
            "A symbolic benchmark grounded in music-theory and musicology concepts",
            "A multi-task design that combines classification and transformation tasks",
            "An automatic evaluation pipeline with canonical prediction fields",
            "An optional explanation layer for qualitative analytical discussion",
        ],
    )


def add_related_work(document: Document) -> None:
    document.add_heading("2. Related Work", level=1)
    document.add_paragraph(
        "This section should position the project against prior work in symbolic music modeling, "
        "music-information retrieval, computational musicology, and the evaluation of language "
        "models on structured reasoning tasks."
    )
    add_bullets(
        document,
        [
            "Symbolic music representation and sequence modeling",
            "Rule-based or theory-informed music-analysis systems",
            "Benchmarks for structured or symbolic reasoning",
            "Computational musicology approaches to analytical categories",
        ],
    )
    document.add_paragraph(
        "Writing prompt: explain clearly what is already known, what existing systems evaluate, "
        "and how the current benchmark differs in scope, task design, or interpretive ambition."
    )


def add_methodology(document: Document) -> None:
    document.add_heading("3. Methodology", level=1)
    document.add_paragraph(
        "This section should explain the benchmark as a full pipeline: task design, data generation, "
        "symbolic representation, prompt construction, prediction schema, and evaluation."
    )

    document.add_heading("3.1 Benchmark Design", level=2)
    add_bullets(
        document,
        [
            "Task family A: label classification",
            "Task family B: note-sequence transformation",
            "Zero-shot benchmark design",
            "Canonical prediction schema with backward compatibility",
        ],
    )

    document.add_heading("3.2 Task Suite", level=2)
    add_numbered(
        document,
        [
            "Interval Identification",
            "Chord Identification",
            "Harmonic Function Classification",
            "Transposition",
            "Melodic Inversion",
            "Retrograde",
            "Rhythmic Augmentation / Diminution",
            "Voice-Leading Violation Detection",
        ],
    )

    document.add_heading("3.3 Symbolic Representations and Tokenizers", level=2)
    document.add_paragraph(
        "Describe the benchmark's symbolic format and explain the role of note-level, MIDI-like, "
        "or REMI-like views if they are used in the final experiments."
    )

    document.add_heading("3.4 Prediction Interface", level=2)
    add_bullets(
        document,
        [
            "Label tasks use prediction_label",
            "Sequence tasks use prediction_notes",
            "prediction_structured is optional for easier downstream inspection",
            "prediction_explanation is optional and reserved for explanatory mode",
        ],
    )

    document.add_heading("3.5 Standard and Explanatory Modes", level=2)
    document.add_paragraph(
        "The benchmark should be described as having a standard benchmark mode and an explanatory "
        "mode. The standard mode is the primary mode and is designed for stable automatic scoring. "
        "The explanatory mode is supplementary and is intended for paper-oriented qualitative analysis."
    )

    document.add_heading("3.6 Current V1 Scope", level=2)
    add_bullets(
        document,
        [
            "Harmonic function is limited to major mode in V1",
            "Voice-leading is simplified to two voices in V1",
            "The harmonic vocabulary is intentionally reduced in V1",
            "The benchmark focuses on symbolic music rather than audio",
        ],
    )
    document.add_paragraph(
        "This subsection should explicitly state that these choices define the current V1 scope "
        "and not the final theoretical boundary of the project."
    )


def add_experimental_setup(document: Document) -> None:
    document.add_heading("4. Experimental Setup", level=1)
    document.add_paragraph(
        "Describe the models, prompting conditions, case generation settings, tokenizer views, "
        "and evaluation scripts used in the study."
    )
    add_bullets(
        document,
        [
            "Model(s) evaluated",
            "Prompt mode(s) used",
            "Number of cases per task",
            "Evaluation scripts and output files",
            "Any qualitative review procedure for explanatory outputs",
        ],
    )


def add_results(document: Document) -> None:
    document.add_heading("5. Results", level=1)
    document.add_paragraph(
        "This section should present the main empirical results. If experiments are still ongoing, "
        "you can use this structure as a placeholder and later replace draft notes with final numbers."
    )
    add_experiment_table(document)
    add_caption(document, "Table 1", "Suggested summary table for main benchmark findings.")

    document.add_heading("5.1 Quantitative Findings", level=2)
    document.add_paragraph(
        "Report the main benchmark metrics here. Separate findings by task family or by model when useful."
    )

    document.add_heading("5.2 Qualitative Findings", level=2)
    document.add_paragraph(
        "Use this subsection to compare formally correct predictions with analytical explanations, "
        "especially where the explanatory mode reveals strengths or limitations that are not obvious "
        "from the strict benchmark metric alone."
    )


def add_discussion(document: Document) -> None:
    document.add_heading("6. Discussion", level=1)
    add_bullets(
        document,
        [
            "Interpret what the benchmark measures successfully",
            "Discuss where symbolic accuracy aligns or fails to align with musicological explanation",
            "Explain what the V1 scope allows and what it leaves out",
            "Relate the results back to the research questions",
        ],
    )


def add_limitations(document: Document) -> None:
    document.add_heading("7. Limitations and Planned Extensions", level=1)
    document.add_paragraph(
        "This section should be explicit and honest. The benchmark is already useful in its current "
        "form, but it should not be presented as complete coverage of tonal analysis or musicology."
    )
    add_bullets(
        document,
        [
            "Current V1 scope is intentionally constrained for reproducibility",
            "Theoretical coverage is not yet final",
            "Planned extensions include broader harmonic-function coverage, richer voice-leading settings, and expanded harmonic vocabulary",
            "Future work may compare benchmark accuracy and explanation quality more systematically",
        ],
    )


def add_conclusion(document: Document) -> None:
    document.add_heading("8. Conclusion", level=1)
    document.add_paragraph(
        "A strong conclusion should restate the problem, summarize the proposed benchmark, highlight "
        "the most important findings, and close with the broader significance of the work for symbolic "
        "music evaluation and computational musicology."
    )


def add_references(document: Document) -> None:
    document.add_heading("References", level=1)
    document.add_paragraph("[Insert references here in your chosen citation style.]")
    document.add_paragraph("[Example placeholder] Author, A. (Year). Title. Venue.")


def add_appendix(document: Document) -> None:
    document.add_heading("Appendix A. Reproducibility Checklist", level=1)
    document.add_paragraph(
        "This appendix is optional but useful if you want the paper to look more complete and systematic."
    )
    add_repro_table(document)
    add_caption(document, "Table A1", "Suggested reproducibility checklist.")


def build_document() -> Document:
    document = Document()
    set_page_layout(document)
    set_metadata(document)
    style_document(document)
    add_footer(document)

    add_title_block(document)
    document.add_page_break()
    add_front_matter(document)
    add_introduction(document)
    add_related_work(document)
    add_methodology(document)
    add_experimental_setup(document)
    add_results(document)
    add_discussion(document)
    add_limitations(document)
    add_conclusion(document)
    add_references(document)
    add_appendix(document)
    return document


def main() -> int:
    document = build_document()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_PATH)
    print(f"wrote: {OUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
